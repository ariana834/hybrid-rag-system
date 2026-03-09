import re
import uuid
from io import BytesIO
from typing import Optional

import docx
from pypdf import PdfReader

from models.document import Document


class DocumentParser:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def parse(self, uploaded_file) -> Optional[Document]:
        try:
            file_name = uploaded_file.name
            file_name_lower = file_name.lower()

            if file_name_lower.endswith(".pdf"):
                text = self._extract_pdf(uploaded_file)
                file_type = "pdf"
            elif file_name_lower.endswith(".docx"):
                text = self._extract_docx(uploaded_file)
                file_type = "docx"
            elif file_name_lower.endswith(".txt"):
                text = self._extract_txt(uploaded_file)
                file_type = "txt"
            else:
                print(f"Unsupported file format: {file_name}")
                return None

            text = self._clean_text(text)
            if not text:
                print(f"No text extracted from file: {file_name}")
                return None

            return Document(
                id=str(uuid.uuid4()),
                filename=file_name,
                file_type=file_type,
                content=text,
                num_characters=len(text),
            )

        except Exception as error:
            print(f"Parser error for {getattr(uploaded_file, 'name', 'unknown file')}: {error}")
            return None

    def _extract_pdf(self, uploaded_file) -> str:
        file_bytes = uploaded_file.getvalue()
        pdf_stream = BytesIO(file_bytes)
        reader = PdfReader(pdf_stream)

        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                pages.append(page_text.strip())

        return "\n\n".join(pages)

    def _extract_docx(self, uploaded_file) -> str:
        file_bytes = uploaded_file.getvalue()
        docx_stream = BytesIO(file_bytes)
        document = docx.Document(docx_stream)

        paragraphs = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        return "\n".join(paragraphs)

    def _extract_txt(self, uploaded_file) -> str:
        file_bytes = uploaded_file.getvalue()

        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue

        return file_bytes.decode("utf-8", errors="ignore")

    def _clean_text(self, text: str) -> str:
        text = text.replace("\r", "\n")
        text = text.replace("\t", " ")
        text = re.sub(r"[ ]{2,}", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" *\n *", "\n", text)
        return text.strip()